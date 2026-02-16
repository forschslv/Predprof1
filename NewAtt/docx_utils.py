from docx import Document
from docx.shared import Cm, Pt
import os


def generate_table_setting_report(orders_data, filename="table_report.docx", period: str = None, grand_total: float = None):
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Заголовок периода, если передан
    if period:
        p_header = document.add_paragraph()
        run_h = p_header.add_run(period)
        run_h.bold = True
        run_h.font.size = Pt(12)
        p_header.paragraph_format.space_after = Pt(6)

    for order in orders_data:
        # Ожидаем структуру: { 'user_name': 'Имя Фамилия', 'dishes': [...], 'total': float }
        name = order.get('user_name', '').strip()
        dishes = order.get('dishes', []) or []
        total = order.get('total', 0.0) or 0.0

        # Имя пользователя
        p_name = document.add_paragraph()
        run = p_name.add_run(name)
        run.bold = True
        run.font.size = Pt(11)

        # Список блюд (каждое в новую строку или через запятую если кратко)
        if dishes:
            # Ограничим длинные списки — покажем по одному на строку
            for d in dishes:
                p_d = document.add_paragraph()
                p_d.paragraph_format.left_indent = Cm(0.5)
                run_d = p_d.add_run(str(d))
                run_d.font.size = Pt(11)

        # Итоговая сумма
        p_total = document.add_paragraph()
        p_total.paragraph_format.space_after = Pt(8)
        run_t = p_total.add_run(f"Итого: {total:.2f} ₽")
        run_t.bold = True
        run_t.font.size = Pt(11)

    os.makedirs("reports", exist_ok=True)
    file_path = os.path.join("reports", filename)
    document.save(file_path)
    # Итог по всем пользователям
    if grand_total is not None:
        p_final = document.add_paragraph()
        run_f = p_final.add_run(f"Итого по всем: {grand_total:.2f} ₽")
        run_f.bold = True
        run_f.font.size = Pt(12)

    document.save(file_path)
    return file_path